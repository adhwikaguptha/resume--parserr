from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import pymongo
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
import spacy
import os
import re
from dotenv import load_dotenv
from bson import ObjectId, binary
import json
import base64
import datetime
import logging
import requests
from pdf2image import convert_from_path
import pytesseract
import tempfile
import io
import asyncio
import aiohttp
from typing import Dict, List, Optional
from werkzeug.utils import secure_filename
import validators
import time

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class JSONEncoder(json.JSONEncoder):
    def default(self, o):
        if isinstance(o, ObjectId):
            return str(o)
        elif isinstance(o, binary.Binary):
            return base64.b64encode(o).decode('utf-8')
        return super().default(o)

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
app = Flask(__name__, static_folder=os.path.join(BASE_DIR, "static"), template_folder=os.path.join(BASE_DIR, "templates"))

# Flask JSON provider
try:
    from flask.json.provider import DefaultJSONProvider
    class CustomJSONProvider(DefaultJSONProvider):
        def dumps(self, obj, **kwargs):
            kwargs.setdefault("cls", JSONEncoder)
            return super().dumps(obj, **kwargs)
        def loads(self, s, **kwargs):
            return super().loads(s, **kwargs)
    app.json_provider_class = CustomJSONProvider
except ImportError:
    app.json_encoder = JSONEncoder

CORS(app)
load_dotenv()

class Config:
    MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/anthropic_resumeparser")
    ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
    UPLOAD_FOLDER = os.path.join(BASE_DIR, "Uploads")
    ALLOWED_MIMETYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    MAX_CONTENT_LENGTH = 5 * 1024 * 1024  # 5 MB

    def __init__(self):
        if not self.ANTHROPIC_API_KEY:
            raise ValueError("ANTHROPIC_API_KEY is not set in .env")

app.config.from_object(Config())

# MongoDB setup
try:
    client = pymongo.MongoClient(app.config["MONGO_URI"], serverSelectionTimeoutMS=5000)
    db = client["anthropic_resumeparser"]
    user_collection = db["users_anthropic"]
    profile_collection = db["user_profile_data"]
    profile_collection.create_index([("username", pymongo.ASCENDING)])
    client.admin.command('ping')
    logger.info("MongoDB connected successfully")
except Exception as e:
    logger.error(f"MongoDB connection failed: {e}")
    raise Exception(f"MongoDB connection failed: {e}")

# spaCy setup
try:
    nlp = spacy.load("en_core_web_lg")
except Exception as e:
    raise Exception(f"spaCy load error: {e}")

# Ensure upload folder exists
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# Constants
STATES = [
    "Alabama", "Alaska", "Arizona", "Arkansas", "California", "Colorado", "Connecticut", "Delaware", "Florida",
    "Georgia", "Hawaii", "Idaho", "Illinois", "Indiana", "Iowa", "Kansas", "Kentucky", "Louisiana", "Maine",
    "Maryland", "Massachusetts", "Michigan", "Minnesota", "Mississippi", "Missouri", "Montana", "Nebraska",
    "Nevada", "New Hampshire", "New Jersey", "New Mexico", "New York", "North Carolina", "North Dakota", "Ohio",
    "Oklahoma", "Oregon", "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota", "Tennessee", "Texas",
    "Utah", "Vermont", "Virginia", "Washington", "West Virginia", "Wisconsin", "Wyoming"
]

SECTION_ALIASES = {
    "career_objective": {"career objective", "objective", "professional summary", "summary", "profile", "about me", "career goals", "personal profile"},
    "education": {"education", "academic background", "academic qualifications", "degrees"},
    "experience": {"experience", "work experience", "professional experience", "employment history", "work history"},
    "skills": {"skills", "technical skills", "soft skills", "languages", "competencies", "abilities"},
    "projects": {"projects", "project experience", "personal projects", "portfolio"},
    "certifications": {"certifications", "certificates", "credentials", "licenses"},
    "achievements": {"achievements", "awards", "honors", "accomplishments"},
    "social_media": {"social media", "online profiles", "links", "contact links"}
}

def normalize_section_name(name: str) -> str:
    name_clean = name.lower().strip(":").strip().replace(" ", "_")
    for standard, aliases in SECTION_ALIASES.items():
        if name_clean in aliases:
            return standard
    return name_clean

def safe_join_list(items: List) -> str:
    return "\n".join(str(i).strip() for i in items if isinstance(i, (str, int, float)) and i not in ["...", Ellipsis])

def ensure_list(obj) -> List:
    if isinstance(obj, list):
        return obj
    elif obj is None:
        return []
    return [obj]

def extract_text_from_pdf(file_path: str) -> str:
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text(layout=True, x_tolerance=2, y_tolerance=2)
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            text += " | ".join(str(cell or "") for cell in row) + "\n"
            if text.strip():
                return text.encode("utf-8", errors="ignore").decode("utf-8").strip()
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
    
    try:
        pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_PATH", r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe")
        images = convert_from_path(file_path, poppler_path=os.getenv("POPPLER_PATH", r"C:\\Program Files\\poppler\\bin"))
        ocr_text = ""
        for img in images:
            ocr_text += pytesseract.image_to_string(img, lang="eng") + "\n"
        return ocr_text.encode("utf-8", errors="ignore").decode("utf-8").strip()
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip() and para.text != "..." and para.text != Ellipsis:
                text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text.append(row_text)
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    text.append(header.text.strip())
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    text.append(footer.text.strip())
        return "\n".join(text)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")

def extract_text_from_file(file_path: str, file_type: str) -> str:
    if file_type == "application/pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def preprocess_text(text: str) -> str:
    if isinstance(text, list):
        text = safe_join_list(text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s.,;@\-/]", "", text)
    return text.strip()

def extract_data_spacy_regex(text: str) -> Dict:
    try:
        doc = nlp(text)
    except Exception as e:
        logger.error(f"spaCy processing error: {e}")
        return {"error": str(e)}
    
    data = {}
    for ent in doc.ents:
        if ent.label_ == "PERSON" and "name" not in data:
            data["name"] = ent.text
        elif ent.label_ == "GPE" and "state" not in data:
            if ent.text in STATES:
                data["state"] = ent.text
    
    if "name" not in data:
        lines = text.splitlines()
        data["name"] = lines[0].strip() if lines else "Unknown"
    
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    phone_pattern = r"(\+?\d{1,3})?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?:\s?(?:x|ext\.?)\s?\d{1,5})?"
    social_media_patterns = {
        "linkedin": r"(?:https?://)?(?:www\.)?linkedin\.com/(?:in|pub|company)/[a-zA-Z0-9-]+/?",
        "github": r"(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9-]+/?",
        "twitter": r"(?:https?://)?(?:www\.)?(?:twitter\.com|x\.com)/[a-zA-Z0-9_]+/?",
        "portfolio": r"(?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}/?(?:portfolio)?(?:/[a-zA-Z0-9-]+)?/?",
    }
    
    email = re.search(email_pattern, text, re.IGNORECASE)
    data["email"] = email.group() if email else None
    if data["email"] and not validators.email(data["email"]):
        logger.warning(f"Invalid email format: {data['email']}")
        data["email"] = None
    
    phone = re.search(phone_pattern, text)
    data["phone"] = phone.group() if phone else None
    
    state_pattern = r"\b(" + "|".join(STATES) + r")\b"
    state = re.search(state_pattern, text, re.IGNORECASE)
    if state and "state" not in data:
        data["state"] = state.group()
    
    data["social_media"] = {}
    for platform, pattern in social_media_patterns.items():
        matches = re.findall(pattern, text, re.IGNORECASE)
        valid_urls = [url for url in matches if validators.url(url) or validators.url("https://" + url)]
        if valid_urls:
            data["social_media"][platform] = valid_urls[0]
    
    lines = text.split("\n")
    current_section = None
    data["sections"] = {}
    section_header_pattern = r"^(?:[A-Z][a-zA-Z\s&]{1,50}|[A-Z\s&]{2,50}):?$"
    
    # Enhanced career objective handling
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(section_header_pattern, line, re.IGNORECASE):
            current_section = normalize_section_name(line)
            data["sections"][current_section] = []
        elif current_section:
            # For career_objective, combine multi-line entries if they form a paragraph
            if current_section == "career_objective" and data["sections"][current_section] and len(line) > 10:
                data["sections"][current_section][-1] += " " + line
            elif len(line) > 5:  # Filter out short or irrelevant lines
                data["sections"][current_section].append(line)
    
    # Validate career objective entries
    if "career_objective" in data["sections"]:
        data["sections"]["career_objective"] = [
            entry for entry in data["sections"]["career_objective"] if len(entry.strip()) > 20
        ]
        if not data["sections"]["career_objective"]:
            del data["sections"]["career_objective"]
            logger.warning("Career objective section empty after validation")
    
    return data

async def extract_data_llm(text: str) -> Dict:
    headers = {
        "x-api-key": app.config['ANTHROPIC_API_KEY'],
        "anthropic-version": "2023-06-01",
        "Content-Type": "application/json"
    }
    
    prompt = (
        "You are a resume parsing expert. Extract information from the provided resume text and return a JSON object with the following structure:\n"
        "{\n"
        "  \"name\": \"string or null\",\n"
        "  \"email\": \"string or null\",\n"
        "  \"phone\": \"string or null\",\n"
        "  \"state\": \"string or null\",\n"
        "  \"social_media\": {\"linkedin\": \"string or null\", \"github\": \"string or null\", \"twitter\": \"string or null\", \"portfolio\": \"string or null\", \"other\": [\"string\", ...]},\n"
        "  \"career_objective\": \"string or null\",\n"
        "  \"education\": [{\"institution\": \"string\", \"degree\": \"string\", \"dates\": \"string\", \"details\": \"string\"}, ...] or null,\n"
        "  \"experience\": [{\"company\": \"string\", \"role\": \"string\", \"dates\": \"string\", \"details\": \"string\"}, ...] or null,\n"
        "  \"skills\": {\"technical_skills\": [\"string\", ...], \"soft_skills\": [\"string\", ...], \"languages\": [\"string\", ...], \"other_skills\": [\"string\", ...]},\n"
        "  \"projects\": [{\"name\": \"string\", \"description\": \"string\", \"dates\": \"string\"}, ...] or null,\n"
        "  \"certifications\": [{\"name\": \"string\", \"issuer\": \"string\", \"date\": \"string\"}, ...] or null,\n"
        "  \"achievements\": [\"string\", ...] or null\n"
        "}\n"
        "Rules:\n"
        "- Return valid JSON only. Wrap the response in ```json\n...\n```.\n"
        "- Extract email as a single valid email address.\n"
        "- Extract phone number in a consistent format (e.g., '+1-123-456-7890').\n"
        "- Extract social media links for LinkedIn, GitHub, Twitter/X, and personal portfolio websites. Include other URLs in 'other'.\n"
        "- Extract career objective or summary as a single string, combining paragraphs or bullet points into a cohesive summary. Set to null if not present or ambiguous.\n"
        "- Categorize skills accurately (e.g., 'Python' as technical, 'Teamwork' as soft, 'Spanish' as language).\n"
        "- Handle missing sections by setting them to null or empty lists/objects.\n"
        "- Parse dates in a consistent format (e.g., 'MM/YYYY - MM/YYYY' or 'Present').\n"
        "- For complex layouts, infer sections based on context or common resume patterns.\n"
        "- If a section is ambiguous, place it under 'other_skills' or 'achievements' as appropriate.\n"
        "Resume text:\n" + text[:8000]
    )
    
    payload = {
        "model": "claude-3-5-sonnet-20240620",
        "temperature": 0.5,
        "max_tokens": 2000,
        "messages": [{"role": "user", "content": prompt}]
    }
    
    async def try_request():
        async with aiohttp.ClientSession() as session:
            for attempt in range(3):
                try:
                    async with session.post(
                        "https://api.anthropic.com/v1/messages",
                        headers=headers,
                        json=payload,
                        timeout=60
                    ) as res:
                        if res.status == 200:
                            out = await res.json()
                            msg_content = out.get("content", [])
                            text_part = ""
                            for msg in msg_content:
                                if msg.get("type") == "text":
                                    text_part += msg.get("text", "")
                            json_match = re.search(r"```json\n(.*?)\n```", text_part, re.DOTALL)
                            if json_match:
                                try:
                                    return json.loads(json_match.group(1))
                                except json.JSONDecodeError as e:
                                    logger.error(f"LLM JSON parsing error: {e}")
                                    return {"error": f"Invalid JSON from LLM: {e}", "raw_content": text_part}
                            return {"error": "No JSON block found in LLM response", "raw_content": text_part}
                        else:
                            logger.warning(f"LLM API error: {res.status} {await res.text()}")
                except Exception as e:
                    logger.warning(f"LLM request attempt {attempt + 1} failed: {e}")
                    if attempt == 2:
                        return {"error": f"LLM API request failed after retries: {e}"}
                    await asyncio.sleep(1)
    
    return await try_request()

def structure_resume_for_storage(spacy_data: Dict, llm_data: Dict) -> Dict:
    result = {
        "schema_version": 1,
        "section_order": ["name", "email", "phone", "state", "social_media", "career_objective", "education", "experience", "skills", "projects", "certifications", "achievements"]
    }
    
    for field in ["name", "email", "phone", "state"]:
        result[field] = llm_data.get(field, spacy_data.get(field))
    
    social_media_struct = {
        "linkedin": None,
        "github": None,
        "twitter": None,
        "portfolio": None,
        "other": []
    }
    
    llm_social_media = llm_data.get("social_media", {})
    if isinstance(llm_social_media, dict):
        for platform in social_media_struct:
            if platform in llm_social_media and llm_social_media[platform]:
                social_media_struct[platform] = llm_social_media[platform]
    
    spacy_social_media = spacy_data.get("social_media", {})
    for platform in social_media_struct:
        if social_media_struct[platform] is None and platform in spacy_social_media:
            social_media_struct[platform] = spacy_social_media[platform]
    
    spacy_social_section = spacy_data.get("sections", {}).get("social_media", [])
    for item in spacy_social_section:
        for platform, pattern in {
            "linkedin": r"linkedin\.com",
            "github": r"github\.com",
            "twitter": r"(?:twitter\.com|x\.com)",
            "portfolio": r"(?:portfolio|[a-zA-Z0-9-]+\.[a-zA-Z]{2,}/)"
        }.items():
            if re.search(pattern, item, re.IGNORECASE):
                if platform not in social_media_struct or social_media_struct[platform] is None:
                    social_media_struct[platform] = item
                else:
                    social_media_struct["other"].append(item)
    
    result["social_media"] = {k: v for k, v in social_media_struct.items() if v}
    
    # Handle career objective
    llm_career_objective = llm_data.get("career_objective")
    spacy_career_objective = spacy_data.get("sections", {}).get("career_objective", [])
    
    if isinstance(llm_career_objective, str) and llm_career_objective.strip():
        result["career_objective"] = llm_career_objective
    elif llm_career_objective is not None:
        result["career_objective"] = ensure_list(llm_career_objective)
    elif spacy_career_objective:
        # Combine spaCy entries into a single string if they form a paragraph
        if len(spacy_career_objective) == 1:
            result["career_objective"] = spacy_career_objective[0]
        else:
            result["career_objective"] = " ".join(spacy_career_objective)
        if len(result["career_objective"]) < 20:
            logger.warning(f"Career objective too short: {result['career_objective']}")
            result["career_objective"] = None
    
    # Handle other sections
    for sec in ["education", "experience", "projects", "certifications", "achievements"]:
        llm_entries = llm_data.get(sec, [])
        spacy_entries = spacy_data.get("sections", {}).get(sec, [])
        if llm_entries:
            result[sec] = ensure_list(llm_entries)
        elif spacy_entries:
            structured = []
            for item in spacy_entries:
                if isinstance(item, dict):
                    structured.append(item)
                elif isinstance(item, str):
                    fields = [f.strip() for f in re.split(r"[;,|-]", item) if f.strip()]
                    structured.append({"details": ", ".join(fields)})
            result[sec] = structured
    
    skills_struct = {
        "technical_skills": [],
        "soft_skills": [],
        "languages": [],
        "other_skills": []
    }
    llm_skills = llm_data.get("skills", {})
    if isinstance(llm_skills, dict):
        for category in skills_struct:
            skills_struct[category] = [str(s).strip() for s in ensure_list(llm_skills.get(category, []))]
    else:
        spacy_skills = spacy_data.get("sections", {}).get("skills", [])
        for skill in spacy_skills:
            skill = str(skill).strip().lower()
            if re.match(r"python|java|javascript|c\+\+|sql|html|css|react|angular|node\.js|docker|aws|git", skill, re.IGNORECASE):
                skills_struct["technical_skills"].append(skill)
            elif re.match(r"spanish|french|german|chinese|japanese|english", skill, re.IGNORECASE):
                skills_struct["languages"].append(skill)
            elif re.match(r"teamwork|communication|leadership|problem\s?solving|adaptability", skill, re.IGNORECASE):
                skills_struct["soft_skills"].append(skill)
            else:
                skills_struct["other_skills"].append(skill)
    
    result["skills"] = {k: v for k, v in skills_struct.items() if v}
    
    return result

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/upload", methods=["POST"])
async def upload_resume():
    file_path = None
    try:
        file = request.files.get("resume")
        if not file or file.filename == "":
            return jsonify({"error": "No file uploaded"}), 400
        
        if file.mimetype not in app.config["ALLOWED_MIMETYPES"]:
            return jsonify({"error": f"Unsupported file type: {file.mimetype}"}), 400
        
        username = request.form.get("username")
        if not username:
            return jsonify({"error": "Username is required"}), 400
        
        file_buffer = file.read()
        if len(file_buffer) > app.config["MAX_CONTENT_LENGTH"]:
            return jsonify({"error": "File size exceeds 5MB limit"}), 400
        file.seek(0)
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(file_path)
        
        raw_text = extract_text_from_file(file_path, file.mimetype)
        raw_text = safe_join_list(raw_text) if isinstance(raw_text, list) else raw_text
        
        spacy_data = extract_data_spacy_regex(raw_text)
        if "error" in spacy_data:
            raise Exception(spacy_data["error"])
        
        llm_data = await extract_data_llm(preprocess_text(raw_text))
        if "error" in llm_data:
            raise Exception(llm_data["error"])
        
        structured = structure_resume_for_storage(spacy_data, llm_data)
        structured.update({
            "username": username,
            "pdfText": raw_text,
            "resumePdf": binary.Binary(file_buffer),
            "created_at": datetime.datetime.utcnow(),
            "updated_at": datetime.datetime.utcnow()
        })
        
        profile_collection.update_one(
            {"username": username},
            {"$set": structured},
            upsert=True
        )
        
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        
        filtered_structured = {k: v for k, v in structured.items() if k != "resumePdf"}
        return jsonify({
            "message": "Resume uploaded and processed successfully",
            "pdfText": raw_text,
            "data": filtered_structured
        })
    
    except Exception as e:
        logger.error(f"Upload error: {e}")
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({"error": f"Failed to upload resume: {str(e)}"}), 500

@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory(app.static_folder, filename)

if __name__ == "__main__":
    app.run(debug=True, port=5000)